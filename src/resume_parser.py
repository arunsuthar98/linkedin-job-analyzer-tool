"""Resume parser — extracts clean text from PDF and DOCX uploads.

Design constraints (from security review):
- Max file size: 5 MB
- Max extracted text: 15,000 characters
- Max pages (PDF): 10
- Encrypted PDFs rejected
- Extraction quality check (warns on scanned/image PDFs)
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Optional

MAX_FILE_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_EXTRACTED_CHARS = 15_000
MAX_PDF_PAGES = 10
MIN_TEXT_CHARS = 100  # below this = likely scanned/empty


@dataclass
class ParseResult:
    text: str
    pages: int
    file_type: str
    warning: Optional[str] = None  # e.g. "low text — may be scanned"


def _clean(text: str) -> str:
    """Remove excessive whitespace while preserving structure."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def parse_pdf(data: bytes) -> ParseResult:
    import pdfplumber

    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"File too large ({len(data)//1024} KB). Max 5 MB.")

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        if pdf.metadata.get("Encrypt"):
            raise ValueError("Encrypted PDFs are not supported. Please remove the password first.")

        pages = pdf.pages[:MAX_PDF_PAGES]
        parts: list[str] = []
        for page in pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)

        text = _clean("\n".join(parts))

    warning = None
    if len(text) < MIN_TEXT_CHARS:
        warning = (
            "Very little text was extracted — your PDF may be scanned or image-based. "
            "For best results, upload a DOCX file or paste your resume text manually."
        )

    return ParseResult(
        text=text[:MAX_EXTRACTED_CHARS],
        pages=len(pages),
        file_type="PDF",
        warning=warning,
    )


def parse_docx(data: bytes) -> ParseResult:
    from docx import Document

    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"File too large ({len(data)//1024} KB). Max 5 MB.")

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = _clean("\n".join(parts))

    warning = None
    if len(text) < MIN_TEXT_CHARS:
        warning = "Very little text was extracted from this DOCX file."

    return ParseResult(
        text=text[:MAX_EXTRACTED_CHARS],
        pages=0,
        file_type="DOCX",
        warning=warning,
    )


def parse_resume(filename: str, data: bytes) -> ParseResult:
    """Dispatch to correct parser based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return parse_pdf(data)
    elif ext in ("docx", "doc"):
        return parse_docx(data)
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Please upload a PDF or DOCX file.")
